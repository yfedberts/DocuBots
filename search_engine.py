import requests
from guppy import hpy
import os
from dotenv import load_dotenv, find_dotenv
from nltk import word_tokenize, sent_tokenize
from docx import Document
import analyzer

class SearchEngine():        

    def doc_to_paraset(self, filename):
        doc = Document(filename)
        fullText = []
        for p in doc.paragraphs:
            if p != '':
                fullText.append(p.text)
        fullText = [ft for ft in fullText if ft]
        return ",".join(fullText)

    def getQuery(self, filename):
        se = SearchEngine()

        doc = Document(filename)
        fullText = []
        for p in doc.paragraphs:
            if p != '':
                fullText.append(p.text)
        fullText = [ft for ft in fullText if ft]
        return(fullText)

        """
        paralist = se.doc_to_paraset(filename)
        Text = sent_tokenize(paralist)
        Text = [t for t in Text if t]
        return Text
        """


    def searchWeb(self, qInput):
        load_dotenv(find_dotenv())

        API_KEY = os.getenv("API_KEY4")
        SEARCH_ENGINE_ID = os.getenv("SEARCH_ENGINE_ID4")
        
        queryInput = qInput
        result_links = []
        result_scores = 0.0

        for q in range(len(queryInput)):
            query = queryInput[q]
            url = f"https://www.googleapis.com/customsearch/v1?key={API_KEY}&cx={SEARCH_ENGINE_ID}&q={query}"
            data = requests.get(url).json()

            results_item = data.get("items")
            if results_item == None:
                continue
            snippets = []
            links = []

            for rs in results_item:
                try:
                    snippet = rs.get("snippet")
                    link = rs.get("link")

                    snippets.append(snippet)
                    links.append(link)
                except:
                    continue

            best_score = 0.0
            best_link = ''

            for i in range(len(snippets)):
                simi = analyzer.query_similarity(links[i], query)
                if(simi > best_score):
                    best_score = simi
                    best_link = links[i]

            result_links.append(best_link)
            result_scores += best_score
        
        return((result_scores/len(result_links)) * 100), result_links

    def get_simi_link(self, d):
            
        se = SearchEngine()
        doc = se.getQuery(d)
        link_result, score = se.searchWeb(doc)
        try:
            if(link_result != '' and score != 0.0):
                return link_result, score
        except:
            return "No similar content found online"

#h = hpy()
a = SearchEngine()
print(a.get_simi_link('docxFiles\sample.docx'))
#print(h.heap())