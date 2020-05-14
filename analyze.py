from nltk import word_tokenize, sent_tokenize, pos_tag
import nltk
from docx import Document
from nltk.corpus import wordnet as wn

def tag_conversion(penntag):
    morphy_tag = {'NN':wn.NOUN, 'JJ':wn.ADJ,
                  'VB':wn.VERB, 'RB':wn.ADV}
    try:
        return morphy_tag[penntag[:2]]
    except:
        return None

def tag_to_synset(word, tag):
    wn_tag = tag_conversion(tag)
    if wn_tag != '':
        return None

    try:
        return wn.synsets(word, wn_tag)[0]
    except:
        return None

def analyze_similarity(sentence1, sentence2):
    sent1 = word_tokenize(sentence1)
    sent2 = word_tokenize(sentence2)

    tags1 = pos_tag(sent1)
    tags2 = pos_tag(sent2)

    doc_synsets = []
    doc_synsets2 = []

    for token, tag in zip(sent1,tags1):
        wordnet_tag = tag_conversion(tag[1])
        if wordnet_tag != None:
            syns = wn.synsets(token, wordnet_tag)
            doc_synsets.append(syns)

    for token, tag in zip(sent2,tags2):
        wordnet_tag = tag_conversion(tag[1])
        if wordnet_tag != None:
            syns = wn.synsets(token, wordnet_tag)
            doc_synsets2.append(syns)

    score, count = 0.0, 0

    #ERROR BELOW
    for ss in doc_synsets:
        best_score = max([ss.path_similarity(ss2) for ss2 in doc_synsets2])

        print(best_score)
    



print(analyze_similarity('I am a big dog for it and you', 'I am a small dog for her and you'))


"""
lib1 = []
lib2 = []

for ss in wn.synsets('car'):
    lib1.append(ss)

for ss in wn.synsets('internet'):
    lib2.append(ss)

print(lib1)
print(lib2)
"""
"""
doc = Document(r'docxFiles\sample.docx')
docxContent  = []
docxContentTokens = []
for p in doc.paragraphs:
    if p.text != '':
        docxContent.append(p.text)
        docxContentTokens.append(word_tokenize(p.text))

docToCompare = Document(r'docxFiles\test.docx')
docxCompareContents = []
docxCompareTokens = []
for p in docToCompare.paragraphs:
    if p.text != '':
        docxCompareContents.append(p.text)
        docxCompareTokens.append(word_tokenize(p.text))

print(docxCompareTokens)
print("="*10 + "\n\n")
print(docxContentTokens)


lib1 = []
lib2 = []

doc_synsets =[]

for w in docxContentTokens:
    lib1.append(nltk.pos_tag(w))

print(lib1)

score, count = 0.0, 0
for w in lib1:
    best_score = max(w.path_similarity(w2) for w2 in lib2)

    if best_score != '':
        score += best_score
        count += 1

score /= count

print(score)

#TODO: Iterate through the tokens array
#Make another array of tokens + synsets
#Compare the two
"""