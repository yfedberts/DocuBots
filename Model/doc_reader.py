from nltk import word_tokenize, sent_tokenize, pos_tag
import nltk
from docx import Document
from nltk.corpus import wordnet as wn
import pandas as pd
import os
from pathlib import Path

#Defining pathings because python's imports sucks
CURR_PATH = os.path.dirname(__file__)
DATA_FOLDER = os.path.relpath("..\\Model\\Data")

#Files pathing
DOC_TEXTS = os.path.join(CURR_PATH, DATA_FOLDER, "doc_texts.csv")

class DocReader():

    def getText(self, filename):
        """
        Return docx files into one set of string to be tokenized later
        """
        if(os.path.exists(DOC_TEXTS)):
            os.remove(DOC_TEXTS)

        df = pd.DataFrame()
        doc = Document(filename)
        fullText = []
        for p in doc.paragraphs:
            if p.text != '':
                fullText.append(p.text)

        df['ParaText'] = fullText
        df.to_csv(DOC_TEXTS)
        return " ".join(fullText)

    def getQuery(self, filename):
        """
        Read the docx files and separate the content into paragraphs
        Return content of each paragraphs in a list to be used for
        Search functions later
        """
        if(os.path.exists(DOC_TEXTS)):
            os.remove(DOC_TEXTS)

        df = pd.DataFrame()
        doc = Document(filename)
        fullText = []
        for p in doc.paragraphs:
            if p != '':
                fullText.append(p.text)
        fullText = [ft for ft in fullText if ft]
        df['ParaText'] = fullText
        df.to_csv(DOC_TEXTS)
        return(fullText)

    def tag_conversion(self, penntag):
        """
        Convert penn tree tag into a WordNet Tag
        """
        morphy_tag = {'NN':wn.NOUN, 'JJ':wn.ADJ,
                    'VB':wn.VERB, 'RB':wn.ADV}
        try:
            return morphy_tag[penntag[:2]]
        except KeyError:
            return None

    def doc_to_synsets(self, doc):
        """
        Getting text from a document file to be tokenized and separated into words,
        WordNet is then used to find synonym sets for each words to be used for comparisons later
        """
        Text = self.getText(doc)
        token = word_tokenize(Text)
        tag = nltk.pos_tag(token)
        wn_tag = [(i[0], self.tag_conversion(i[1])) for i in tag]
        output = [wn.synsets(i, z)[0] for i, z in wn_tag if len(wn.synsets(i, z))>0]
        return output

    def str_to_synsets(self, s1):
        """
        Instead of reading from a docx files, it reads in strings and
        returns a the sets of synonyms.
        """
        token = word_tokenize(s1)
        tag = nltk.pos_tag(token)
        wn_tag = [(i[0], self.tag_conversion(i[1])) for i in tag]
        output = [wn.synsets(i, z)[0] for i, z in wn_tag if len(wn.synsets(i, z))>0]
        return output

    def similarity_ratio(self, s1,s2):
        """
        S1 = Synonym sets of input of documents/queries
        S2 = Synonym sets of text from docs/queries to be compared with

        Compares synonym sets of S1 and S2 and returns the highest comparison score
        """
        simis = []
        simlist = []

        for s in s1:
            simlist = [s.path_similarity(ss) for ss in s2 if s.path_similarity(ss) is not None]
            if not simlist:
                continue

            max_score = max(simlist)
            simis.append(max_score)
        try:
            output = sum(simis)/len(simis)
            return output
        except:
            output = 0.0
            return output

    def doc_similarity(self, d1, d2):
        """
        Function used to compare simiarity between 2 docx files
        """
        synsets1 = self.doc_to_synsets(d1)
        synsets2 = self.doc_to_synsets(d2)

        return ((self.similarity_ratio(synsets1,synsets2) + self.similarity_ratio(synsets2, synsets1)) /2)

    def query_similarity(self, s1, s2):
        """
        Function used to compare queries, can also be used for sentences/string comparison
        """
        synsets1 = self.str_to_synsets(s1)
        synsets2 = self.str_to_synsets(s2)

        return((self.similarity_ratio(synsets1, synsets2) + self.similarity_ratio(synsets2, synsets1)) / 2)

    def quick_compare(self, s1, s2):
        return((self.similarity_ratio(s1,s2) + self.similarity_ratio(s2,s1))/2)